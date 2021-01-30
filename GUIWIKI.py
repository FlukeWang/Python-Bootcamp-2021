# GUIWIKI.py
import wikipedia

# python to doc
from docx import Document
def Wiki(keyword, lang='en'):

    wikipedia.set_lang(lang)
    # Article Summary
    # summary is สรุป so the content is brief
    # data = wikipedia.summary(keyword) use data 2 instead

    # Article (page + content)
    datafull = wikipedia.page(keyword)
    datafull = datafull.content

    doc = Document()  # create word file in python
    doc.add_heading(keyword, 0)  # 0 heading size

    doc.add_paragraph(datafull)
    doc.save(keyword + '.docx')

    print('File Created')


# Chang language to TH
# wikipedia.set_lang('th') ; I # by myself cuz I want Eng
from tkinter import *
from tkinter import ttk  # file ย่อย ใน tkinter
from tkinter import messagebox  # to show function error when can't search


GUI = Tk()  # work with mainloop
GUI.title('WIKI By Fluke Wang')
GUI.geometry('400x300')

# config

FONT1 = ('Ariel', 20)

# Label

L = ttk.Label(GUI, text='Enter your Keywords', font=FONT1)
L.pack(pady=5)

# Search Entry
v_search = StringVar()  # Keyword Box
E1 = ttk.Entry(GUI, textvariable=v_search, font=FONT1, width=40)
E1.pack(pady=5)  # ขยายช่องแนว y

# Search Button


def search():
    keyword = v_search.get()  # .get() = ดึงข้อมูลเข้ามา , สำหรับ StringVar เท่านั้น
    try:  # try to find results if ok pass it
        language = v_radio.get()  # th/en/zh from radio botton downside
        Wiki(keyword, language)
        messagebox.showinfo('Keyword correct', 'File created')
    except:  # if don't have result pop up this
        # ERROR is title and Please try again is message
        messagebox.showwarning('Keyword Error!', 'Please Try Again')
    # print(wikipedia.search(keyword))
    # result = (wikipedia.summary(keyword))
    # print(result)
    # sharp all because we already create Wiki function


B1 = ttk.Button(GUI, text='Search', command=search)
B1.pack(padx=5, pady=5)


# Change languagte Menu
F1 = Frame(GUI)  # whiteboard and put radio button on it
F1.pack(pady=10)  # made F1(frame) with include radiobotton have gap

v_radio = StringVar()  # Data box (lang)

RB1 = ttk.Radiobutton(F1, text='ภาษาไทย', variable=v_radio, value='th')
RB2 = ttk.Radiobutton(F1, text='English', variable=v_radio, value='en')
RB3 = ttk.Radiobutton(F1, text='中文', variable=v_radio, value='zh')
RB2.invoke()  # Set english as default

# .pack mean verticle but we need make it look good by change to horizontal
# RB1.pack()
# RB2.pack()
# RB3.pack()

RB1.grid(row=0, column=0, )  # .grid = horizontal
RB2.grid(row=0, column=1, )
RB3.grid(row=0, column=2, )


GUI.mainloop()
