#test-docx.py
from docx import Document #file docx.py class document():
import wikipedia # file wikipedoa.py def summary():


def Wiki(keyword,lang = 'th') : 

    wikipedia.set_lang(lang)
    # Article Summary
    data = wikipedia.summary(keyword) # summary is สรุป so the content is brief

    # Article (page + content)
    datafull = wikipedia.page(keyword)
    datafull = datafull.content

    doc = Document() # create word file in python 
    doc.add_heading(keyword,0) # 0 heading size 

    doc.add_paragraph(datafull)
    doc.save(keyword + '.docx')

    print('Created')

Wiki('Japan','th')
Wiki('TSMC','zh')
Wiki('Benz','fr')
